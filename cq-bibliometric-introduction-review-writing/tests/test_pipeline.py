from __future__ import annotations

import json
import io
import os
import sys
import tempfile
import unittest
from pathlib import Path
from contextlib import redirect_stdout
from unittest.mock import patch

SKILL = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(SKILL / "scripts"))

from analysis import analyze, keyword_list, network_sensitivity_rows, nmf_access_gate
from advanced_analysis import burst_tables, document_clusters, language_tokens, modeling_text, network_metrics, nmf_topics, preprocessing_audit
import credentials
import install_skill
import review_pipeline
import theory_library
from common import canonical_record, deduplicate, ensure_task, file_sha256, load_jsonl, read_csv_optional, read_json, write_json, write_jsonl
from evidence import build_evidence, citation_minimum, citation_trigger_errors, claim_ids_in_text, clean_apa_reference, clean_introduction_text, punctuation_errors, sync_references, sync_review_references, topic_citation_quotas, validate, write_introduction, writing_eligible_records
from deliverables import COPYRIGHT, clean_audit_markers, export_deliverables, markdown_body
from extract import extract_documents
from focus import focus_records
from ingest import ingest
from review_pipeline import main as pipeline_main
from science import prisma_report, quality_appraisal
from fourth_round import citation_age_analysis, kmeans_incremental_diagnostic, knowledge_flow, strategic_map
from semantic import compile_semantic, embedding_dry_run, prepare as prepare_semantic, reconcile_semantic, validate_semantic
from meta_analysis import compile_meta, validate_meta
from platform_support import path_issues
from writing_audit import _atomic_units, audit_writing
from publication import compile_publication, prepare_publication, validate_publication
from gap_design import compile_design, compile_gaps, prepare_design, prepare_gaps, validate_design, validate_gaps
from sources import CachedClient, oa_download_inventory, openalex_to_record, reconstruct_abstract, run_search_plan, validate_search_plan
from theory_library import (compile_recommendations, ensure_library, import_library,
                            ingest_theories, library_status, promote_theories,
                            safe_task_theory_phase, search_theories, set_task_support,
                            theory_home, validate_recommendations, verify_theories)
from workflow_v5 import apply_corpus_policy, build_search_strategy, metadata_coverage, write_brief


class FakeResponse:
    def __init__(self, status, payload=None):
        self.status_code = status; self._payload = payload or {}; self.text = json.dumps(self._payload); self.headers = {}; self.url = "https://example.test"
    def json(self): return self._payload


class WorkflowTests(unittest.TestCase):
    def test_introduction_marker_cleanup_preserves_paragraph_boundaries(self):
        source = (
            "[role:social-importance] 第一段证据 [cite:R1234567890abcd]。\n\n"
            "[role:gap] 第二段缺口 [cite:Rabcdef12345678]。\n"
        )
        for cleaner in (clean_introduction_text, clean_audit_markers):
            cleaned = cleaner(source)
            self.assertEqual(len([p for p in cleaned.split("\n\n") if p.strip()]), 2)
            self.assertNotIn("[role:", cleaned)
            self.assertNotIn("[cite:", cleaned)
        html_body, _ = markdown_body(clean_audit_markers(source))
        self.assertEqual(html_body.count("<p>"), 2)

    def test_windows_path_preflight_and_managed_copy_installer(self):
        issues = path_issues(Path("C:/reviews/CON/task"))
        self.assertTrue(any(x["code"] == "windows-reserved-name" for x in issues))
        with tempfile.TemporaryDirectory() as td:
            source = Path(td) / "source"; source.mkdir(); (source / "SKILL.md").write_text("---\nname: demo\n---\n", encoding="utf-8")
            target = Path(td) / "entry"
            with patch.object(install_skill, "SKILL", source):
                installed = install_skill.install_one("codex", target, "copy")
                self.assertEqual(installed["mode"], "copy"); self.assertEqual(install_skill.inspect_one("codex", target)["status"], "ready")

    def test_credential_backend_failure_falls_back_to_session_without_leak(self):
        credentials.SESSION_CREDENTIALS.clear()

    def test_missing_gui_falls_back_to_non_echoing_terminal_prompt(self):
        with patch("credentials.platform.system", return_value="Linux"), patch("credentials._tk_dialog", side_effect=RuntimeError("no display")), patch("credentials._terminal_prompt", return_value="temporary-secret") as terminal:
            self.assertEqual(credentials.prompt_value("key", "title", hidden=True, input_mode="auto"), "temporary-secret")
            terminal.assert_called_once_with("key", True)
        with patch("credentials.platform.system", return_value="Linux"), patch("credentials._keyring_set", side_effect=RuntimeError("locked")):
            storage = credentials.store_credential("OPENALEX_API_KEY", "not-for-output")
        self.assertEqual(storage, "session-only"); self.assertEqual(credentials.credential_value("OPENALEX_API_KEY"), "not-for-output")
        credentials.SESSION_CREDENTIALS.clear()

    def test_language_preprocessing_is_unicode_aware_and_audited(self):
        tokens = language_tokens("Ｉｎｆｏｒｍａｔｉｏｎｏｖｅｒｌｏａｄ 信息超载影响冲动购买")
        self.assertIn("informationoverload", tokens); self.assertTrue(any("信息" in x or "超载" in x for x in tokens))
        record = canonical_record({"title": "信息超载", "abstract": "消费者冲动购买"})
        audit = preprocessing_audit([record]); self.assertTrue(audit.iloc[0]["detected_cjk"]); self.assertGreater(audit.iloc[0]["tokens"], 0)

    def test_network_null_model_is_reported_as_signal_only(self):
        pd = __import__("pandas")
        edges = pd.DataFrame([{"a": "a", "b": "b", "w": 3}, {"a": "b", "b": "c", "w": 3}, {"a": "a", "b": "c", "w": 3}])
        rows = network_sensitivity_rows("demo", edges, "a", "b", "w", "fractional")
        self.assertTrue(rows); self.assertEqual(rows[0]["interpretation"], "robustness-signal-only"); self.assertEqual(rows[0]["null_runs"], 10)

    def test_conditional_meta_analysis_and_insufficient_fallback(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); records = [canonical_record({"title": f"Study {i}"}) for i in range(10)]
            write_jsonl(root / "02_corpus/corpus.jsonl", records)
            effects = [{"effect_id": f"E{i:04d}", "record_id": record["record_id"], "sample_id": f"sample-{i}", "outcome": "purchase", "metric": "r", "estimate": .10 + i / 100, "n": 150 + i * 10, "anchor": "abstract", "extraction_status": "completed", "subgroup": "A" if i < 5 else "B", "numeric_moderator": i} for i, record in enumerate(records)]
            write_jsonl(root / "05_evidence/meta/effect_sizes.jsonl", effects)
            report = compile_meta(root); self.assertEqual(report["status"], "ready"); self.assertEqual(report["syntheses"][0]["publication_bias_status"], "exploratory-egger")
            self.assertTrue(validate_meta(root)["valid"])
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "One"}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_jsonl(root / "05_evidence/meta/effect_sizes.jsonl", [{"effect_id": "E1", "record_id": record["record_id"], "sample_id": "s1", "outcome": "x", "metric": "r", "estimate": .2, "n": 100, "anchor": "abstract", "extraction_status": "completed"}])
            self.assertEqual(compile_meta(root)["status"], "skipped-insufficient-data")

    def test_semantic_cache_and_selective_reconciliation(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence", "abstract": "Association evidence."}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            prepared = prepare_semantic(root); task = json.loads(next((root / "05_evidence/semantic/batches").glob("*.json")).read_text(encoding="utf-8"))["tasks"][0]
            extraction = task["template"]; extraction["host_review_status"] = "completed"; extraction["anchors"] = ["abstract"]
            write_json(root / f"05_evidence/semantic/extractions/{record['record_id']}.json", extraction)
            self.assertEqual(prepare_semantic(root)["pending_records"], 0)
            compile_semantic(root); self.assertEqual(reconcile_semantic(root)["tasks"], 0)

    def test_sentence_audit_requires_support_not_just_a_citation(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence"}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            (root / "06_review/review_draft.md").write_text(f"# Review\n\n研究表明两者相关 [cite:{record['record_id']}]。\n", encoding="utf-8")
            first = audit_writing(root, "review"); self.assertFalse(first["valid"]); self.assertTrue(any("pending" in x for x in first["errors"]))
            ledger = load_jsonl(root / "05_evidence/prose_claim_ledger.jsonl"); ledger[0].update({"support_status": "supported", "design_fit": "associational-language", "evidence_level": "abstract", "claim_scope": "association", "audit_notes": "The cited abstract directly supports the stated association.", "independent_samples": 1})
            write_jsonl(root / "05_evidence/prose_claim_ledger.jsonl", ledger)
            self.assertTrue(audit_writing(root, "review")["valid"])

    def test_kmeans_is_diagnostic_and_gated(self):
        pd = __import__("pandas")
        assignments = pd.DataFrame([{"record_id": f"R{i}", "topic_id": i % 2 + 1} for i in range(12)])
        clusters = pd.DataFrame([{"record_id": f"R{i}", "cluster_id": i % 3 + 1} for i in range(12)])
        diagnostics = pd.DataFrame([{"k": 3, "silhouette": .30, "stability": .90, "min_cluster_size": 4}])
        _, status = kmeans_incremental_diagnostic(assignments, clusters, diagnostics, {"selected_k": 3})
        self.assertEqual(status["role"], "heterogeneity-diagnostic-only"); self.assertTrue(status["writing_allowed"])

    def test_strategic_map_never_calls_weak_theme_emerging(self):
        pd = __import__("pandas")
        records = [canonical_record({"title": f"Document {i}", "abstract": "topic evidence"}) for i in range(50)]
        topics = pd.DataFrame([{"topic_id": i, "documents": 10, "top_terms": f"term{i}a; term{i}b"} for i in range(1, 6)])
        assignments = pd.DataFrame([{"record_id": r["record_id"], "topic_id": i % 5 + 1} for i, r in enumerate(records)])
        edges = pd.DataFrame([{"term_a": "term1a", "term_b": "term1b", "cooccurrence": 4}, {"term_a": "term1a", "term_b": "term2a", "cooccurrence": 2}])
        frame, status = strategic_map(records, topics, assignments, edges, "on")
        self.assertFalse(any("emerg" in x for x in frame["quadrant"].astype(str))); self.assertIn(status["status"], {"ready", "exploratory"})

    def test_citation_age_and_half_life(self):
        pd = __import__("pandas"); records = []
        for i in range(20):
            records.append(canonical_record({"title": f"Paper {i}", "year": 2020, "references": [{"year": y} for y in range(2010, 2016)]}))
        assignments = pd.DataFrame([{"record_id": r["record_id"], "topic_id": i % 4 + 1} for i, r in enumerate(records)])
        instances, summary, status = citation_age_analysis(records, assignments)
        self.assertEqual(status["status"], "ready"); self.assertEqual(len(instances), 120); self.assertIn("synchronous_half_life", summary.columns)

    def test_knowledge_flow_gates_sparse_network(self):
        pd = __import__("pandas")
        records = [canonical_record({"title": f"Paper {i}", "year": 2020 + i % 2}) for i in range(10)]
        assignments = pd.DataFrame([{"record_id": r["record_id"], "topic_id": i % 2 + 1} for i, r in enumerate(records)])
        edges = pd.DataFrame([{"source_id": records[i]["record_id"], "target_id": records[(i+1)%10]["record_id"], "weight": 1} for i in range(10)])
        _, _, status = knowledge_flow(edges, assignments, records)
        self.assertEqual(status["status"], "skipped-insufficient-data"); self.assertFalse(status["writing_allowed"])

    def test_knowledge_flow_uses_permutation_null_gate(self):
        pd = __import__("pandas"); records = [canonical_record({"title": f"Paper {i}", "year": 2020 + i % 2}) for i in range(30)]
        assignments = pd.DataFrame([{"record_id": r["record_id"], "topic_id": 1 if i < 15 else 2} for i, r in enumerate(records)])
        edges = []
        for citing in records[15:]:
            for cited in records[:4]: edges.append({"source_id": citing["record_id"], "target_id": cited["record_id"], "weight": 1})
        _, flows, status = knowledge_flow(pd.DataFrame(edges), assignments, records)
        self.assertEqual(status["null_runs"], 100); self.assertIn("null_z", flows.columns); self.assertTrue(status["writing_allowed"])

    def test_semantic_workflow_and_embedding_dry_run(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Survey evidence", "abstract": "A survey reports an association.", "year": 2024})
            write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            prepared = prepare_semantic(root, 10); self.assertFalse(prepared["embedding_used"])
            extraction = json.loads(next((root / "05_evidence/semantic/batches").glob("*.json")).read_text(encoding="utf-8"))["tasks"][0]["template"]
            extraction.update({"host_review_status": "completed", "design": "survey", "anchors": ["abstract"], "relations": [{"subject": "X", "predicate": "is associated with", "object": "Y", "direction": "positive", "status": "support", "anchor": "abstract"}]})
            (root / f"05_evidence/semantic/extractions/{record['record_id']}.json").write_text(json.dumps(extraction), encoding="utf-8")
            self.assertEqual(compile_semantic(root)["relations"], 1); self.assertTrue(validate_semantic(root)["valid"])
            dry = embedding_dry_run(root); self.assertFalse(dry["download_started"]); self.assertEqual(dry["status"], "interface-only")

    def test_base_requirements_have_no_embedding_stack(self):
        requirements = (SKILL / "scripts/requirements.txt").read_text(encoding="utf-8").lower()
        for forbidden in ("torch", "sentence-transformers", "flagembedding", "transformers"):
            self.assertNotIn(forbidden, requirements)

    def test_validation_blocks_gated_kmeans_claim(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Known", "year": 2020})
            write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_jsonl(root / "05_evidence/claim_ledger.jsonl", [{"claim_id": "C0001", "claim_type": "bibliometric", "record_ids": [record["record_id"]], "anchors": ["04_analysis/tables/kmeans_incremental_cross.csv"]}])
            (root / "04_analysis/advanced_module_status.json").write_text(json.dumps({"kmeans": {"writing_allowed": False}}), encoding="utf-8")
            report = validate(root); self.assertFalse(report["valid"]); self.assertTrue(any("KMeans".lower() in x.lower() or "kmeans" in x.lower() for x in report["errors"]))
    def test_credential_status_masks_and_warns_environment_override(self):
        with patch.dict(os.environ, {"OPENALEX_API_KEY": "environment-secret"}, clear=False), patch("credentials._stored_value", return_value=("stored-secret", "keychain")):
            report = credentials.credential_status(); item = report["credentials"]["OPENALEX_API_KEY"]
            self.assertTrue(item["environment_overrides_secure_store"]); self.assertEqual(item["source"], "environment")
            self.assertNotIn("environment-secret", json.dumps(report)); self.assertEqual(len(item["fingerprint"]), 8)

    def test_credential_test_classifies_missing_without_secret(self):
        with patch("credentials.credential_value", return_value=""):
            report = credentials.test_credentials()
        self.assertFalse(report["ok"]); self.assertEqual(report["category"], "missing"); self.assertTrue(report["secrets_exposed"] is False)

    def test_manifest_schema_migrates_without_touching_corpus(self):
        from common import update_manifest
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); corpus = root / "02_corpus/corpus.jsonl"; corpus.write_text('{"title":"keep"}\n', encoding="utf-8")
            before = corpus.read_bytes(); update_manifest(root, "status")
            manifest = json.loads((root / "manifest.json").read_text(encoding="utf-8"))
            self.assertEqual(manifest["skill_schema_version"], 11); self.assertEqual(before, corpus.read_bytes())

    def test_publication_variant_rewrites_workflow_language_without_upgrading_partial_claim(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence", "authors": ["Li, A"], "year": 2024}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_jsonl(root / "05_evidence/reference_registry.jsonl", [{"record_id": record["record_id"], "apa": "Li, A. (2024). Evidence.", "doi": ""}])
            source = root / "06_review/review_draft.md"; source.write_text(f"# Review\n\n现有摘要证据不足以证明该机制 (Li, 2024) [cite:{record['record_id']}]。\n", encoding="utf-8")
            first = audit_writing(root, "review")
            ledger = load_jsonl(root / "05_evidence/prose_claim_ledger.jsonl"); ledger[-1].update({"support_status": "partial", "design_fit": "causal-language-narrowed", "evidence_level": "abstract", "claim_scope": "mechanism proposition", "audit_notes": "Indirect support only.", "independent_samples": 1}); write_jsonl(root / "05_evidence/prose_claim_ledger.jsonl", ledger)
            self.assertTrue(audit_writing(root, "review")["valid"])
            prepared = prepare_publication(root, "review"); self.assertGreater(prepared["queue_items"], 0)
            publication = root / "06_review/review_publication_audit.md"
            publication.write_text(f"# Review\n\n本研究提出该机制可能成立 (Li, 2024) [cite:{record['record_id']}]。\n", encoding="utf-8")
            compile_publication(root, "review")
            publication_row = next(row for row in load_jsonl(root / "05_evidence/prose_claim_ledger.jsonl") if row.get("variant") == "publication")
            write_jsonl(root / "05_evidence/publication_support_reviews.jsonl", [{"document": "review", "sentence_id": publication_row["sentence_id"], "support_status": "partial", "design_fit": "causal-language-narrowed", "evidence_level": "abstract", "claim_scope": "testable mechanism proposition", "audit_notes": "Wording remains explicitly propositional.", "independent_samples": 1}])
            report = validate_publication(root, "review")
            self.assertTrue(report["valid"], report["errors"]); self.assertFalse(report["banned_terms"])
            self.assertTrue((root / "06_review/review_publication.md").exists())
            self.assertEqual(prepare_publication(root, "review")["queue_items"], 0)

    def test_publication_atomic_audit_blocks_heterogeneous_end_cluster(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); records = [canonical_record({"title": f"Evidence {i}"}) for i in range(8)]; write_jsonl(root / "02_corpus/corpus.jsonl", records)
            ids = ",".join(r["record_id"] for r in records)
            source = root / "06_review/review_publication_audit.md"
            source.write_text(f"# Review\n\n历史重演、公共史与文化表演表明服装和动作能够组织体验；数字叙事则连接地点与历史信息 [cite:{ids}]。\n", encoding="utf-8")
            report = audit_writing(root, "review", variant="publication")
            self.assertFalse(report["valid"]); self.assertGreater(report["atomic_repair_queue_items"], 0)
            queue = load_jsonl(root / "05_evidence/citation_attribution_queue.jsonl")
            self.assertTrue(any("atomic-claim-lacks-exact-citation" in row["reason_codes"] for row in queue))

    def test_atomic_split_ignores_semicolons_inside_apa_and_audit_groups(self):
        sentence = "第一机制得到支持 (Li, 2024; Wang, 2025) [cite:R1234567890abcd,Rabcdef12345678]；第二机制仍待检验。"
        units = _atomic_units(sentence)
        self.assertEqual(len(units), 2)
        self.assertIn("Li, 2024; Wang, 2025", units[0])
        self.assertTrue(units[0].endswith("；"))

    def test_prisma_seed_recall_and_quality_appraisal(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Known experimental paper", "abstract": "An experiment with N=120 participants and regression method.", "year": 2024})
            (root / "00_plan/search_plan.json").write_text(json.dumps({"seed_papers": [{"title": "Known experimental paper"}]}), encoding="utf-8")
            report = prisma_report(root, "screening", [record]); self.assertEqual(report["seed_recall"]["recall"], 1.0)
            quality = quality_appraisal([record])[0]; self.assertEqual(quality["study_design"], "experiment"); self.assertTrue(quality["reports_sample_or_material"])

    def test_environment_credential_overrides_secure_store(self):
        with patch.dict(os.environ, {"OPENALEX_API_KEY": "environment-secret"}, clear=False):
            with patch("credentials._macos_get") as keychain:
                self.assertEqual(credentials.credential_value("OPENALEX_API_KEY"), "environment-secret")
                keychain.assert_not_called()

    def test_dialog_configuration_stores_without_returning_values(self):
        with patch("credentials.prompt_value", side_effect=["openalex-secret-value", "researcher@example.org"]), patch("credentials.store_credential") as store, patch("credentials.show_message"):
            report = credentials.configure_dialog()
        self.assertTrue(report["configured"]); self.assertNotIn("openalex-secret-value", json.dumps(report))
        store.assert_any_call("OPENALEX_API_KEY", "openalex-secret-value")
        store.assert_any_call("UNPAYWALL_EMAIL", "researcher@example.org")

    def test_search_opens_dialog_when_key_missing(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); plan = {
                "title_or_idea": "中文题目", "title_zh": "中文题目", "title_en": "English title", "source_language": "zh", "translation_status": "completed", "approved": True,
                "concepts": [{"zh": ["中文"], "en": ["English"]}],
                "queries": [{"id": "Q-ZH", "family": "core", "language": "zh", "query": "中文查询"}, {"id": "Q-EN", "family": "core", "language": "en", "query": "English query"}],
            }
            (root / "00_plan/search_plan.json").write_text(json.dumps(plan, ensure_ascii=False), encoding="utf-8")
            with patch("review_pipeline.credential_value", return_value=""), patch("review_pipeline.configure_dialog", return_value={"configured": True}) as dialog, patch("review_pipeline.run_search_plan", return_value=([], {"queries": []})), patch("review_pipeline.export_corpus"):
                code = pipeline_main(["search", "--task", str(root), "--enrich-limit", "0"])
            self.assertEqual(code, 0); dialog.assert_called_once()

    def test_chinese_init_requires_host_translation(self):
        with tempfile.TemporaryDirectory() as td:
            task = Path(td) / "task"
            self.assertEqual(pipeline_main(["init", "--task", str(task), "--title", "数字化转型对企业创新的影响"]), 0)
            plan = json.loads((task / "00_plan/search_plan.json").read_text(encoding="utf-8"))
            self.assertEqual(plan["source_language"], "zh"); self.assertEqual(plan["title_zh"], "数字化转型对企业创新的影响")
            self.assertEqual(plan["title_en"], ""); self.assertEqual(plan["translation_status"], "required")
            report = validate_search_plan(plan); self.assertFalse(report["valid"]); self.assertTrue(any("英文 title_en" in x for x in report["errors"]))

    def test_completed_bilingual_plan_passes(self):
        plan = {
            "title_or_idea": "数字化转型对企业创新的影响", "title_zh": "数字化转型对企业创新的影响",
            "title_en": "The impact of digital transformation on firm innovation", "source_language": "zh", "translation_status": "completed",
            "concepts": [{"name": "数字化转型", "zh": ["数字化转型"], "en": ["digital transformation"]}],
            "queries": [
                {"id": "Q01-ZH", "family": "core", "language": "zh", "query": "数字化转型 企业创新"},
                {"id": "Q01-EN", "family": "core", "language": "en", "query": "digital transformation firm innovation"},
            ],
        }
        report = validate_search_plan(plan); self.assertTrue(report["valid"], report["errors"])

    def test_search_blocks_incomplete_chinese_plan_before_api(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); plan_path = root / "00_plan/search_plan.json"
            plan_path.write_text(json.dumps({"title_or_idea": "中文题目", "title_zh": "中文题目", "title_en": "", "source_language": "zh", "translation_status": "required", "approved": True, "concepts": [], "queries": [{"id": "Q01-ZH", "family": "core", "language": "zh", "query": "中文题目"}]}, ensure_ascii=False), encoding="utf-8")
            with self.assertRaisesRegex(RuntimeError, "双语校验失败"): run_search_plan(root, plan_path)

    def test_english_plan_does_not_require_translation(self):
        plan = {"title_or_idea": "Digital transformation and innovation", "title_en": "Digital transformation and innovation", "source_language": "en", "translation_status": "not_required", "queries": [{"id": "Q01-EN", "family": "core", "language": "en", "query": "digital transformation innovation"}]}
        self.assertTrue(validate_search_plan(plan)["valid"])

    def test_reconstruct_openalex_abstract(self):
        self.assertEqual(reconstruct_abstract({"world": [1], "hello": [0]}), "hello world")

    def test_openalex_normalization(self):
        row = openalex_to_record({"id": "https://openalex.org/W1", "doi": "https://doi.org/10.1/ABC", "title": "Test", "publication_year": 2024, "abstract_inverted_index": {"A": [0], "test": [1]}, "cited_by_count": 3, "counts_by_year": [{"year": 2024, "cited_by_count": 2}], "authorships": [{"author": {"display_name": "Li"}}], "keywords": [{"display_name": "AI"}]}, "Q01")
        self.assertEqual(row["abstract"], "A test"); self.assertEqual(row["ids"]["doi"], "10.1/abc"); self.assertEqual(row["query_ids"], ["Q01"])
        self.assertEqual(row["citation_counts_by_year"], {"2024": 2})

    def test_dedup_merges_sources(self):
        a = canonical_record({"title": "Same study", "year": 2022, "doi": "10.1234/X", "abstract": "short", "citation_counts": {"wos": 4}}, "wos")
        b = canonical_record({"title": "Same study", "year": 2022, "doi": "https://doi.org/10.1234/x", "abstract": "a much longer abstract", "citation_counts": {"openalex": 7}}, "openalex")
        rows, audit = deduplicate([a, b])
        self.assertEqual(len(rows), 1); self.assertEqual(len(audit), 1); self.assertEqual(rows[0]["citation_counts"], {"wos": 4, "openalex": 7}); self.assertEqual(rows[0]["abstract"], "a much longer abstract")

    def test_jsonl_loader_preserves_unicode_line_separator_inside_metadata(self):
        with tempfile.TemporaryDirectory() as td:
            path = Path(td) / "records.jsonl"
            write_jsonl(path, [{"title": "line one\u2028line two"}, {"title": "next record"}])
            rows = load_jsonl(path)
            self.assertEqual(len(rows), 2); self.assertEqual(rows[0]["title"], "line one\u2028line two")

    def test_focus_separates_core_theory_and_excluded_records(self):
        direct = canonical_record({"title": "Information overload and purchase decisions in livestream commerce", "abstract": "Consumers experience information overload while live shopping.", "query_ids": ["Q01-EN"]})
        theory = canonical_record({"title": "Information overload in consumer purchase decision making", "abstract": "Consumer purchase intention declines under information overload.", "query_ids": ["Q07-EN"]})
        health = canonical_record({"title": "Information overload in healthcare livestream education", "abstract": "Patient decision making during COVID-19.", "query_ids": ["Q01-EN"]})
        wos_direct = canonical_record({"title": "Live shopping information overload and consumer buying", "abstract": "Purchase intention in live commerce.", "ids": {"wos": "WOS:1"}})
        core, supplements, excluded, report = focus_records([direct, theory, health, wos_direct])
        self.assertEqual(len(core), 2); self.assertEqual(len(supplements), 1); self.assertEqual(len(excluded), 1)
        self.assertEqual(report["core_records"], 2); self.assertEqual(core[0]["inclusion"]["status"], "included_core")

    def test_keyword_modeling_excludes_automatic_openalex_topics(self):
        record = {"keywords": ["livestream shopping"], "topics": [{"name": "impulse physics"}]}
        self.assertEqual(keyword_list(record), ["livestream shopping"])
        self.assertIn("impulse physics", keyword_list(record, include_topics=True))

    def test_modeling_text_excludes_derived_metadata(self):
        record = {"title": "Livestream trust", "abstract": "Consumers form purchase intentions.", "keywords": ["impulse physics"], "topics": [{"name": "mathematics"}]}
        text = modeling_text(record)
        self.assertIn("Livestream trust", text); self.assertNotIn("physics", text); self.assertNotIn("mathematics", text)

    def test_nmf_and_kmeans_choose_k_independently(self):
        records = []
        groups = [("trust streamer credibility", "purchase intention trust expertise"), ("interface overload information", "cognitive load information overload"), ("comments social presence", "bullet comments social interaction")]
        for group, (title, abstract) in enumerate(groups):
            for i in range(8): records.append(canonical_record({"title": f"{title} {i}", "abstract": f"{abstract} mechanism sample {group} variant{i}", "year": 2015+i%5}))
        topics, assignments, diagnostics, meta = nmf_topics(records, 3, 5)
        clusters, _, cluster_diagnostics, cluster_meta = document_clusters(records, 2, 6)
        self.assertEqual(len(assignments), 24); self.assertEqual(len(clusters), 24)
        self.assertIn(meta["selected_k"], {3, 4, 5}); self.assertIn(cluster_meta["selected_k"], {2, 3, 4, 5, 6})
        self.assertIn("coherence", diagnostics.columns); self.assertIn("silhouette", cluster_diagnostics.columns)
        self.assertIn("bootstrap_stability", diagnostics.columns); self.assertIn("preprocessing_stability", diagnostics.columns); self.assertIn("topic_entropy", assignments.columns)

    def test_kleinberg_burst_has_interval_and_support(self):
        records = []
        for year in range(2015, 2023):
            for i in range(5):
                keyword = ["algorithmic commerce"] if year >= 2020 and i < 4 else ["baseline"]
                records.append(canonical_record({"title": f"Paper {year} {i}", "abstract": "algorithmic commerce recommendation" if keyword[0].startswith("algorithmic") else "general consumer behavior", "year": year, "keywords": keyword}, "wos"))
        assignments = __import__("pandas").DataFrame([{"record_id": r["record_id"], "topic_id": 1} for r in records])
        keywords, _, _, meta = burst_tables(records, keyword_list, assignments, 3)
        self.assertEqual(meta["method"], "kleinberg-two-state"); self.assertFalse(keywords.empty)
        self.assertTrue(any(keywords["start_year"] >= 2020)); self.assertTrue(keywords["supporting_documents"].str.contains("R").any())

    def test_network_structural_holes_include_support(self):
        pd = __import__("pandas")
        edges = pd.DataFrame([
            {"a": "a", "b": "b", "w": 3}, {"a": "b", "b": "c", "w": 3}, {"a": "c", "b": "a", "w": 3},
            {"a": "x", "b": "y", "w": 3}, {"a": "y", "b": "z", "w": 3}, {"a": "z", "b": "x", "w": 3}, {"a": "c", "b": "x", "w": 1},
        ])
        support = {node: {canonical_record({"title": node})["record_id"]} for node in "abcxyz"}
        summary, metrics, _ = network_metrics("test", edges, "a", "b", "w", support)
        self.assertEqual(summary.iloc[0]["components"], 1); self.assertIn("constraint", metrics.columns)
        self.assertTrue(metrics["supporting_documents"].str.contains("R").all())

    def test_dynamic_citation_minimum(self):
        self.assertEqual(citation_minimum(64), 39); self.assertEqual(citation_minimum(100), 60)
        self.assertEqual(citation_minimum(188), 113); self.assertEqual(citation_minimum(400), 120)

    def test_semantic_denominator_uses_planned_batches_not_previous_targets(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); records = [canonical_record({"title": f"P{i}"}) for i in range(6)]; write_jsonl(root / "02_corpus/corpus.jsonl", records)
            write_json(root / "manifest.json", {"corpus_mode": "all"})
            tasks = [{"record_id": r["record_id"]} for r in records]; write_json(root / "05_evidence/semantic/batches/batch-001.json", {"tasks": tasks})
            for i, record in enumerate(records): write_json(root / f"05_evidence/semantic/extractions/{record['record_id']}.json", {"record_id": record["record_id"], "host_review_status": "completed", "relevance": "direct" if i < 4 else "peripheral"})
            write_jsonl(root / "05_evidence/citation_coverage_targets.jsonl", [{"record_id": "Robsolete000000"}])
            self.assertEqual(len(writing_eligible_records(root, records)), 4)

    def test_write_introduction_requires_approval_and_cleans_ids(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td) / "task"); source = Path(td) / "audit.md"
            record = canonical_record({"title": "Evidence", "year": 2024}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_jsonl(root / "05_evidence/reference_registry.jsonl", [{"record_id": record["record_id"], "apa": "Li, A. (2024). Evidence.", "doi": ""}])
            source.write_text(f"# Audit\n\n第一段证据 [cite:{record['record_id']}]。\n\n第二段。", encoding="utf-8")
            with self.assertRaisesRegex(RuntimeError, "尚未获用户确认"): write_introduction(root, source)
            (root / "06_review/outline.md").write_text("> status: approved", encoding="utf-8")
            report = write_introduction(root, source); self.assertEqual(report["status"], "written")
            clean = (root / "06_review/ssci_introduction.md").read_text(encoding="utf-8")
            self.assertNotIn("# Audit", clean); self.assertNotIn(record["record_id"], clean); self.assertIn("## 参考文献", clean)
            self.assertIn("Li, A. (2024). Evidence.", clean)

    def test_claim_range_expansion(self):
        self.assertEqual(claim_ids_in_text("[C0001–C0004]"), {"C0001", "C0002", "C0003", "C0004"})

    def _reference_sync_fixture(self, root: Path):
        records = [canonical_record({"title": "First evidence", "year": 2020}), canonical_record({"title": "Second evidence", "year": 2021})]
        write_jsonl(root / "02_corpus/corpus.jsonl", records)
        write_jsonl(root / "05_evidence/claim_ledger.jsonl", [{"claim_id": "C0001", "claim_type": "abstract", "record_ids": [records[0]["record_id"]], "anchors": ["card"]}])
        write_jsonl(root / "05_evidence/reference_registry.jsonl", [{"record_id": r["record_id"], "apa": f"Author. ({r['year']}). {r['title']}.", "doi": ""} for r in records])
        (root / "06_review/outline.md").write_text("# 提纲\n\n> status: approved\n\n## 证据综合", encoding="utf-8")
        body = f"# Review\n\n## 证据综合\n\n综合论断 [C0001] 及第二项证据 [{records[1]['record_id']}]。"
        (root / "06_review/review_draft.md").write_text(body, encoding="utf-8")
        return records

    def test_sync_references_exactly_matches_body(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); records = self._reference_sync_fixture(root)
            report = sync_review_references(root); self.assertEqual(report["embedded_references"], 2)
            text = (root / "06_review/review_draft.md").read_text(encoding="utf-8")
            self.assertEqual(text.count("<!-- record:"), 2)
            validation = validate(root); self.assertTrue(validation["valid"], validation["errors"])
            broken = text.replace(f"<!-- record:{records[1]['record_id']} -->", "")
            (root / "06_review/review_draft.md").write_text(broken, encoding="utf-8")
            validation = validate(root); self.assertFalse(validation["valid"]); self.assertTrue(any("未列入文末" in x for x in validation["errors"]))

    def test_validation_rejects_posthoc_evidence_heading(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); self._reference_sync_fixture(root)
            draft = root / "06_review/review_draft.md"
            draft.write_text(draft.read_text(encoding="utf-8") + "\n\n### 扩展主题证据\n\n补充。", encoding="utf-8")
            sync_review_references(root)
            report = validate(root); self.assertFalse(report["valid"]); self.assertTrue(any("后置引用填充章节" in x for x in report["errors"]))

    def test_ris_import_and_end_to_end(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td) / "task"); ris = Path(td) / "sample.ris"
            blocks = []
            for i in range(1, 9):
                blocks.append(f"TY  - JOUR\nTI  - Topic evidence study {i}\nAU  - Author, A{i}\nPY  - {2015+i}\nAB  - This study examines mechanism context outcome number {i}.\nKW  - mechanism; context; outcome{i%3}\nDO  - 10.1000/test{i}\nTC  - {i}\nER  -\n")
            ris.write_text("\n".join(blocks), encoding="utf-8")
            records, report = ingest([ris]); self.assertEqual(len(records), 8); self.assertEqual(report["unique_records"], 8)
            write_jsonl(root / "02_corpus/corpus.jsonl", records)
            result = analyze(root, records); self.assertEqual(result["record_count"], 8)
            evidence = build_evidence(root); self.assertEqual(evidence["records"], 8)
            self.assertTrue(validate(root)["valid"])

    def test_text_extraction_and_reference_list(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td) / "task"); doc = Path(td) / "paper.md"
            doc.write_text("# Introduction\nFinding.\n# References\n1. Smith J. Example title. 2020. doi:10.1234/test\n", encoding="utf-8")
            report = extract_documents(root, [doc]); self.assertEqual(report[0]["status"], "ok"); self.assertEqual(report[0]["reference_candidates"], 1)
            extracted = next((root / "03_fulltext/extracted").glob("*.md")).read_text(encoding="utf-8")
            self.assertIn("line:1", extracted)

    def test_wos_bibtex_and_csv_imports(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            wos = td / "savedrecs.txt"
            wos.write_text("FN Clarivate\nVR 1.0\nPT J\nAU Wang, X\nTI A WoS title\nAB First abstract line\n   continued abstract line\nPY 2021\nDI 10.2000/wos\nER\nEF\n", encoding="utf-8")
            bib = td / "refs.bib"
            bib.write_text('@article{x,\n title={A multiline\n BibTeX title},\n author={Li, A and Smith, B},\n year={2020},\n doi={10.2000/bib}\n}\n', encoding="utf-8")
            csv_path = td / "cnki.csv"
            csv_path.write_text("标题,作者,年份,摘要,DOI\n中文题名,张三;李四,2019,中文摘要,10.2000/cnki\n", encoding="utf-8-sig")
            records, report = ingest([wos, bib, csv_path])
            self.assertEqual(report["unique_records"], 3)
            wos_record = next(r for r in records if (r["ids"] or {}).get("doi") == "10.2000/wos")
            self.assertIn("continued abstract line", wos_record["abstract"])
            self.assertTrue(any("multiline BibTeX title" in r["title"] for r in records))

    def test_wos_ris_preserves_wos_identifier_and_citation_count(self):
        with tempfile.TemporaryDirectory() as td:
            ris = Path(td) / "wos.ris"
            ris.write_text("TY  - JOUR\nTI  - A WoS RIS paper\nAU  - Li, A\nPY  - 2024\nDO  - 10.2000/wos-ris\nAN  - WOS:001234567890123\nN1  - Times Cited in Web of Science Core Collection:  42\nER  -\n", encoding="utf-8")
            records, _ = ingest([ris])
            self.assertEqual(len(records), 1)
            self.assertEqual(records[0]["ids"]["wos"], "WOS:001234567890123")
            self.assertEqual(records[0]["citation_counts"]["wos"], 42)

    def test_pdf_and_docx_anchors(self):
        try:
            import fitz
            from docx import Document
        except ImportError:
            self.skipTest("optional PDF/DOCX dependencies are unavailable")
        with tempfile.TemporaryDirectory() as td:
            td = Path(td); root = ensure_task(td / "task")
            pdf = td / "paper.pdf"; document = fitz.open(); page = document.new_page(); page.insert_text((72, 72), "Test paper title\nIntroduction\nA supported finding."); document.save(pdf); document.close()
            docx = td / "paper.docx"; word = Document(); word.add_heading("Test paper title", 0); word.add_paragraph("A Word finding."); word.save(docx)
            report = extract_documents(root, [pdf, docx]); self.assertTrue(all(x["status"] == "ok" for x in report))
            texts = "\n".join(p.read_text(encoding="utf-8") for p in (root / "03_fulltext/extracted").glob("*.md"))
            self.assertIn("page:1", texts); self.assertIn("paragraph:1", texts)

    def test_validation_rejects_unknown_record(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Known", "year": 2020})
            write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_jsonl(root / "05_evidence/claim_ledger.jsonl", [{"claim_id": "C0001", "claim_type": "abstract", "record_ids": ["Rmissing"], "anchors": []}])
            report = validate(root); self.assertFalse(report["valid"]); self.assertTrue(any("不存在" in e for e in report["errors"]))

    def test_cached_retry(self):
        with tempfile.TemporaryDirectory() as td:
            client = CachedClient(Path(td), retries=2)
            client.session.get = unittest.mock.Mock(side_effect=[FakeResponse(429), FakeResponse(200, {"ok": True})])
            with patch("sources.time.sleep", return_value=None): value = client.get_json("https://example.test")
            self.assertTrue(value["ok"]); self.assertEqual(client.session.get.call_count, 2)

    def test_oa_inventory_separates_paid_content(self):
        records = [
            canonical_record({"title": "A", "oa": {"pdf_url": "https://repo.test/a.pdf"}}),
            canonical_record({"title": "B", "fulltext": {"content_url": "https://content.openalex.org/works/W1.pdf"}}),
        ]
        report = oa_download_inventory(records, 10)
        self.assertEqual(report["direct_oa_available"], 1); self.assertEqual(report["openalex_paid_content_available"], 1); self.assertEqual(report["maximum_estimated_cost_usd"], 0.01)

    def test_strategy_only_generates_wos_without_credentials(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td))
            write_json(root / "00_plan/search_plan.json", {"title_or_idea": "信息超载", "title_zh": "信息超载", "title_en": "Information overload", "source_language": "zh", "translation_status": "completed", "approved": True, "concepts": [{"name": "overload", "zh": ["信息超载"], "en": ["information overload"], "exclude": []}, {"name": "consumer", "zh": ["消费者"], "en": ["consumer decision"], "exclude": []}], "queries": [{"id": "Q1-ZH", "family": "core", "language": "zh", "query": "信息超载 消费者"}, {"id": "Q1-EN", "family": "core", "language": "en", "query": "information overload consumer decision"}]})
            with patch("credentials.credential_value", side_effect=AssertionError("credentials must not be read")):
                report = build_search_strategy(root, "strategy-only")
            self.assertFalse(report["openalex_required"]); self.assertIn("TS=(", (root / "00_plan/wos_search_query.txt").read_text(encoding="utf-8"))

    def test_large_corpus_topic_quota_is_proportional(self):
        quotas = topic_citation_quotas({1: 1820, 2: 599, 3: 690, 4: 473}, 3582)
        self.assertEqual(sum(quotas.values()), 302); self.assertLess(quotas[1], 728)

    def test_ris_extended_metadata_and_coverage(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td); ris = td / "sample.ris"
            ris.write_text("TY  - JOUR\nTI  - Metadata paper\nAU  - Li, A\nY1  - 2024/05/03\nDO  - 10.1000/meta\nTC  - 12\nCR  - Smith J, 2018, Journal, DOI 10.1000/ref\nER  -\n", encoding="utf-8")
            records, _ = ingest([ris]); record = records[0]
            self.assertEqual(record["year"], 2024); self.assertTrue(record["publication_date"].startswith("2024-05")); self.assertEqual(record["citation_counts"]["ris"], 12); self.assertEqual(record["reference_metadata"][0]["year"], 2018)
            root = ensure_task(td / "task"); coverage = metadata_coverage(root, records); self.assertEqual(coverage["fields"]["reference_years"], 1.0)

    def test_all_policy_zero_exclusions(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); records = [canonical_record({"title": f"P{i}"}) for i in range(3)]; write_jsonl(root / "02_corpus/corpus.jsonl", records)
            report = apply_corpus_policy(root, "all"); self.assertEqual(report["excluded"], 0)
            self.assertTrue(all(r["inclusion"]["status"] == "included_all" for r in load_jsonl(root / "02_corpus/corpus.jsonl")))

    def test_deliverables_single_html_browser_word_and_ris(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence", "authors": ["Li, A"], "year": 2024, "doi": "10.1000/e"}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_jsonl(root / "05_evidence/claim_ledger.jsonl", [{"claim_id": "C0001", "record_ids": [record["record_id"]]}])
            write_jsonl(root / "05_evidence/reference_registry.jsonl", [{"record_id": record["record_id"], "apa": "Li, A. (2024). Evidence. https://doi.org/10.1000/e", "doi": "10.1000/e"}])
            (root / "06_review/review_draft.md").write_text(f"# Review\n\nSupported claim [C0001]\n<!-- evidence: {record['record_id']} -->\n", encoding="utf-8")
            report = export_deliverables(root, "review", draft=True)
            html_text = Path(report["html"]).read_text(encoding="utf-8")
            self.assertEqual(report["cited_records"], 1); self.assertNotIn("docx", report); self.assertFalse((root / "06_review/deliverables/review.docx").exists())
            self.assertIn("function exportWord()", html_text); self.assertIn("0x04034b50", html_text); self.assertIn(COPYRIGHT, html_text)
            self.assertIn("DO  - 10.1000/e", Path(report["ris"]).read_text(encoding="utf-8"))

    def test_optional_csv_reader_and_empty_gap_prepare_degrade_cleanly(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); semantic = root / "05_evidence/semantic"; semantic.mkdir(parents=True, exist_ok=True)
            variants = {
                "missing.csv": None,
                "zero.csv": b"",
                "bom.csv": b"\xef\xbb\xbf  \n",
                "header.csv": b"record_id,subject\n",
            }
            for name, payload in variants.items():
                path = semantic / name
                if payload is not None: path.write_bytes(payload)
                frame = read_csv_optional(path, ["record_id", "subject"])
                self.assertTrue(frame.empty); self.assertEqual(list(frame.columns)[:2], ["record_id", "subject"])
            (semantic / "counter_and_null_evidence.csv").write_bytes(b"\xef\xbb\xbf")
            report = prepare_gaps(root)
            self.assertEqual(report["candidates"], 0); self.assertIn("不代表没有研究缺口", report["note"])

    def test_section_audit_builds_semantic_repair_queue(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Relevant evidence", "abstract": "A is associated with B."}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_json(root / f"05_evidence/semantic/extractions/{record['record_id']}.json", {"record_id": record["record_id"], "evidence_level": "abstract", "design": "survey", "contexts": ["LARP"], "variables": {"outcomes": ["travel intention"]}, "relations": [{"direction": "positive"}], "relevance": "direct", "host_review_status": "completed"})
            section = root / "06_review/sections/s1.md"; section.write_text(f"研究表明该关系存在。段末集中引文 [cite:{record['record_id']}]。\n", encoding="utf-8")
            report = audit_writing(root, "review", section, "section-1")
            self.assertFalse(report["valid"]); self.assertGreater(report["repair_queue_items"], 0)
            queue = [x for x in load_jsonl(root / "05_evidence/citation_repair_queue.jsonl") if x["scope"] == "section-1"]
            self.assertTrue(queue); self.assertEqual(queue[0]["candidate_evidence"][0]["design"], "survey")

    def test_only_exact_duplicate_citation_fragment_is_auto_removed(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); a = canonical_record({"title": "A"}); b = canonical_record({"title": "B"}); write_jsonl(root / "02_corpus/corpus.jsonl", [a, b])
            source = root / "06_review/review_draft.md"
            source.write_text(f"# Review\n\n过渡 [cite:{a['record_id']}]。(A, 2024) [cite:{a['record_id']}]\n\n另一段。(B, 2024) [cite:{b['record_id']}]\n", encoding="utf-8")
            report = audit_writing(root, "review")
            self.assertEqual(report["automatic_duplicate_citation_repairs"], 1)
            self.assertEqual(report["automatic_citation_position_repairs"], 1)
            repaired = source.read_text(encoding="utf-8")
            self.assertNotIn("(A, 2024)", repaired); self.assertIn("(B, 2024)", repaired)
            self.assertNotRegex(repaired, r"。\(B, 2024\)")

    def test_nmf_claim_and_structure_gates_are_independent(self):
        pd = __import__("pandas")
        topics = pd.DataFrame([{"topic_id": 1, "documents": 5}, {"topic_id": 2, "documents": 5}])
        unstable = {"selected_converged": True, "documents": 10, "selected_k": 2, "strong_bibliometric_claims_allowed": False}
        auto = nmf_access_gate(unstable, topics, 10, "auto")
        self.assertFalse(auto["claim_allowed"]); self.assertTrue(auto["structure_allowed"]); self.assertEqual(auto["status"], "exploratory-structure")
        self.assertFalse(nmf_access_gate(unstable, topics, 10, "strict")["structure_allowed"])
        self.assertFalse(nmf_access_gate({**unstable, "selected_converged": False}, topics, 10, "auto")["structure_allowed"])

    def test_final_export_blocks_failed_or_stale_validation(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence"}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            source = root / "06_review/review_draft.md"; source.write_text(f"# Review\n\nClaim [cite:{record['record_id']}].\n", encoding="utf-8")
            write_json(root / "06_review/review_approval.json", {"status": "approved", "source_sha256": file_sha256(source)})
            with self.assertRaisesRegex(RuntimeError, "总体验证未通过"):
                export_deliverables(root, "review")
            write_json(root / "07_logs/validation_report.json", {"valid": True, "source_hashes": {"review": "stale"}})
            with self.assertRaisesRegex(RuntimeError, "正文在最近验证后发生变更"):
                export_deliverables(root, "review")

    def test_ten_introduction_paragraphs_survive_html_export(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td))
            paragraphs = [f"[role:gap] 第{i}段保留换行。" for i in range(1, 11)]
            (root / "06_review/ssci_introduction_audit.md").write_text("\n\n".join(paragraphs) + "\n", encoding="utf-8")
            report = export_deliverables(root, "introduction", draft=True)
            html_text = Path(report["html"]).read_text(encoding="utf-8")
            self.assertEqual(html_text.count("<p>"), 10)

    def test_review_approval_gates_introduction_brief(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); source = root / "06_review/review_draft.md"
            source.write_text("# Review\n\n概念性过渡。\n", encoding="utf-8")
            self.assertEqual(pipeline_main(["writing-brief", "--task", str(root), "--document", "introduction", "--skip"]), 1)
            self.assertTrue(audit_writing(root, "review")["valid"])
            export_deliverables(root, "review", draft=True)
            self.assertEqual(pipeline_main(["document-approval", "--task", str(root), "--document", "review", "--status", "approved"]), 0)
            self.assertEqual(pipeline_main(["writing-brief", "--task", str(root), "--document", "introduction", "--skip"]), 0)
            brief = json.loads((root / "06_review/introduction_brief.json").read_text(encoding="utf-8"))
            self.assertEqual(brief["review_approval_sha256"], file_sha256(source))

    def test_clean_audit_markers_removes_groups_without_punctuation_debris(self):
        text = "结论（作者，2024） [cite:R1234567890abcd,Rabcdef12345678]。\n另一句 [R1234567890abcd; Rabcdef12345678]。\n旧标记（R1234567890abcd，page:1）。\n主题 \\[C0001–C0004\\]。"
        clean = clean_audit_markers(text)
        self.assertNotRegex(clean, r"\[\s*[,;，；]+\s*\]")
        self.assertNotIn("R1234567890abcd", clean)
        self.assertNotIn("C0001", clean)
        self.assertFalse(punctuation_errors(clean, "clean"))

    def test_evidence_trigger_requires_nearby_traceable_citation(self):
        unsupported = citation_trigger_errors("一些研究指出该机制稳定。后文没有来源。", [], "综述")
        supported = citation_trigger_errors("一些研究指出该机制稳定 [cite:R1234567890abcd]。", [], "综述")
        self.assertTrue(unsupported); self.assertFalse(supported)

    def test_punctuation_audit_blocks_known_artifacts(self):
        defects = punctuation_errors("异常（；）；另一个，；以及；；。", "正文")
        self.assertGreaterEqual(len(defects), 3)

    def test_apa_cleanup_removes_empty_venue_period(self):
        dirty = "Author. (2024). Title. . https://doi.org/10.1000/example"
        clean = clean_apa_reference(dirty)
        self.assertEqual(clean, "Author. (2024). Title. https://doi.org/10.1000/example")
        self.assertTrue(punctuation_errors(dirty, "reference")); self.assertFalse(punctuation_errors(clean, "reference"))

    def test_introduction_reference_sync_includes_social_sources(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence", "year": 2024}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_jsonl(root / "05_evidence/reference_registry.jsonl", [{"record_id": record["record_id"], "apa": "Li, A. (2024). Evidence.", "doi": ""}])
            write_jsonl(root / "06_review/social_context_sources.jsonl", [{"source_id": "SCTX-001", "source_name": "统计机构", "title": "网络零售数据", "publication_date": "2025-01-01", "url": "https://example.test", "retrieved_at": "2026-07-12"}])
            audit = root / "06_review/ssci_introduction_audit.md"; audit.write_text(f"正文（Li, 2024） [cite:{record['record_id']},SCTX-001]。", encoding="utf-8")
            report = sync_references(root, "introduction")
            text = audit.read_text(encoding="utf-8")
            clean = (root / "06_review/ssci_introduction.md").read_text(encoding="utf-8")
            self.assertEqual(report["embedded_references"], 2); self.assertIn("<!-- source:SCTX-001 -->", text); self.assertIn("统计机构", text)
            self.assertNotIn("<!--", clean); self.assertIn("参考文献", clean)

    def test_gap_audit_demotes_distribution_method_and_incomparable_contradiction(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence", "abstract": "Evidence"}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            base = root / "05_evidence/gaps"; base.mkdir(parents=True, exist_ok=True)
            write_jsonl(base / "gap_ledger.jsonl", [
                {"gap_id": "G1", "title": "尚未使用新方法", "level": "A", "failure_type": "underused-method", "status": "validated"},
                {"gap_id": "G2", "title": "研究结论矛盾", "level": "A", "failure_type": "contradiction", "known_findings": [{"claim": "mixed", "record_ids": [record["record_id"]]}], "current_explanation": ["theory"], "failure_point": "方向不一致", "failure_evidence": [record["record_id"]], "knowledge_consequence": "cannot infer direction", "repair_strategy": "harmonize measures", "research_questions": ["when"], "design_requirements": ["comparable design"], "comparability": {"constructs": "unchecked", "outcomes": "checked", "designs": "checked"}}
            ])
            report = compile_gaps(root); rows = load_jsonl(base / "gap_ledger.jsonl")
            self.assertEqual(report["status"], "opportunity-only"); self.assertTrue(all(x["level"] == "C" for x in rows))
            checked = validate_gaps(root); self.assertTrue(checked["valid"]); self.assertEqual(checked["status"], "opportunity-only")

    def test_valid_explanatory_gap_and_dynamic_method_matching(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Mechanism evidence", "abstract": "Mechanism"}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            base = root / "05_evidence/gaps"; base.mkdir(parents=True, exist_ok=True)
            write_jsonl(base / "gap_ledger.jsonl", [{"gap_id": "GAP-A1", "title": "竞争机制", "level": "A", "failure_type": "competing-explanations", "target_question": "which pathway dominates", "known_findings": [{"claim": "opposing pathways", "record_ids": [record["record_id"]]}], "current_explanation": ["stress", "cue use"], "failure_point": "net effect cannot distinguish pathways", "failure_evidence": [record["record_id"]], "counterevidence": [], "comparability": {"constructs": "checked", "outcomes": "checked", "samples": "checked", "contexts": "checked", "time": "checked", "designs": "checked"}, "knowledge_consequence": "direction remains unexplained", "practical_importance": "consumer welfare", "repair_strategy": "manipulate overload and observe pathways", "research_questions": ["Which pathway dominates?"], "discriminating_predictions": ["opposite outcomes"], "design_requirements": ["randomized manipulation"], "inferential_goals": ["causal", "mechanism"], "confidence": "medium", "evidence_level": "abstract", "status": "validated"}])
            self.assertEqual(compile_gaps(root)["status"], "validated")
            (root / "06_review").mkdir(parents=True, exist_ok=True); write_json(root / "06_review/review_brief.json", {"method_preferences": ["ann"], "method_constraints": []})
            prepare_design(root); compile_design(root); design = validate_design(root)
            ids = {x["method_id"] for x in load_jsonl(root / "05_evidence/design/method_recommendations.jsonl")}
            self.assertIn("experiment", ids); self.assertNotIn("machine-learning", ids); self.assertTrue(any("ann" in x for x in design["warnings"]))

    def test_mechanism_and_context_gaps_require_special_audits(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence", "abstract": "Evidence"}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            base = root / "05_evidence/gaps"; base.mkdir(parents=True, exist_ok=True)
            common = {"known_findings": [{"claim": "supported", "record_ids": [record["record_id"]]}], "current_explanation": ["existing theory"], "failure_point": "unresolved", "failure_evidence": [record["record_id"]], "knowledge_consequence": "explanation remains ambiguous", "repair_strategy": "discriminating design", "research_questions": ["why or when"], "design_requirements": ["temporally ordered evidence"], "status": "validated"}
            write_jsonl(base / "gap_ledger.jsonl", [
                {**common, "gap_id": "GM", "title": "mechanism", "level": "A", "failure_type": "mechanism-missing"},
                {**common, "gap_id": "GB", "title": "context", "level": "B", "failure_type": "external-validity"},
            ])
            compile_gaps(root); rows = {x["gap_id"]: x for x in load_jsonl(base / "gap_ledger.jsonl")}
            self.assertEqual(rows["GM"]["level"], "C"); self.assertEqual(rows["GB"]["level"], "C")
            rows["GM"].update({"level": "A", "status": "validated", "mechanism_audit": {"base_relation_supported": True, "synonyms_checked": True, "statistical_mediation_only": "false"}})
            rows["GB"].update({"level": "B", "status": "validated", "boundary_assumption": "the original theory assumes stable, low-pressure exposure"})
            write_jsonl(base / "gap_ledger.jsonl", list(rows.values()))
            report = compile_gaps(root); self.assertEqual(report["levels"].get("A"), 1); self.assertEqual(report["levels"].get("B"), 1)

    def test_method_matching_changes_with_inferential_goal(self):
        with tempfile.TemporaryDirectory() as td:
            root = ensure_task(Path(td)); record = canonical_record({"title": "Evidence"}); write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            base = root / "05_evidence/gaps"; base.mkdir(parents=True, exist_ok=True)
            rows = []
            for idx, (goal, failure) in enumerate((("prediction", "causal-identification"), ("configuration", "competing-explanations"), ("qualitative", "construct-mismatch")), 1):
                rows.append({"gap_id": f"G{idx}", "title": goal, "level": "B" if idx == 1 else "A", "failure_type": failure, "target_question": goal, "known_findings": [{"claim": "known", "record_ids": [record["record_id"]]}], "current_explanation": ["theory"], "failure_point": "unresolved", "failure_evidence": [record["record_id"]], "counterevidence": [], "comparability": {"constructs": "checked", "outcomes": "checked", "designs": "checked"}, "knowledge_consequence": "limits inference", "practical_importance": "decision quality", "repair_strategy": "fit design", "research_questions": [goal], "discriminating_predictions": ["different observable pattern"], "design_requirements": ["fit data"], "inferential_goals": [goal], "confidence": "medium", "evidence_level": "abstract", "status": "validated"})
            write_jsonl(base / "gap_ledger.jsonl", rows); compile_gaps(root); prepare_design(root); compile_design(root)
            methods = load_jsonl(root / "05_evidence/design/method_recommendations.jsonl")
            by_gap = {gap: {x["method_id"] for x in methods if x["gap_id"] == gap} for gap in ("G1", "G2", "G3")}
            self.assertIn("machine-learning", by_gap["G1"]); self.assertIn("fsqca", by_gap["G2"]); self.assertIn("qualitative-interview", by_gap["G3"])

    def test_method_template_has_no_fixed_three_method_default(self):
        template = (SKILL / "assets/theory_model_package.template.md").read_text(encoding="utf-8").lower()
        self.assertNotIn("pls-sem", template); self.assertNotIn("fsqca", template); self.assertNotRegex(template, r"\bann\b")
        catalog = json.loads((SKILL / "assets/method_catalog.json").read_text(encoding="utf-8"))
        self.assertGreaterEqual(len(catalog), 15); self.assertTrue(all(x.get("cannot_answer") and x.get("assumptions") for x in catalog))

    def test_theory_library_is_independent_empty_and_has_no_builtin_content(self):
        with tempfile.TemporaryDirectory() as td:
            home = ensure_library(Path(td) / "theories")
            report = library_status(home)
            self.assertEqual(report["cards"], 0); self.assertEqual(report["builtin_theories"], 0)
            self.assertFalse(list((SKILL / "assets").glob("theory_cards*.jsonl")))

    def test_theory_home_uses_platform_data_directory_or_override(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td)
            with patch.dict(os.environ, {"CQ_THEORY_HOME": str(base / "override")}):
                self.assertEqual(theory_home(), (base / "override").resolve())
            with patch.dict(os.environ, {"APPDATA": str(base / "roaming")}, clear=False), patch.object(theory_library.platform, "system", return_value="Windows"):
                os.environ.pop("CQ_THEORY_HOME", None); self.assertEqual(theory_home(), (base / "roaming/CQ-BIRW/theories").resolve())
            with patch.dict(os.environ, {"XDG_DATA_HOME": str(base / "xdg")}, clear=False), patch.object(theory_library.platform, "system", return_value="Linux"):
                os.environ.pop("CQ_THEORY_HOME", None); self.assertEqual(theory_home(), (base / "xdg/cq-birw/theories").resolve())

    def test_theory_html_ingest_requires_host_verification_and_confirmation(self):
        with tempfile.TemporaryDirectory() as td:
            base = Path(td); home = base / "library"; source = base / "2025-01-01 研究理论介绍 | 自我决定理论（Self-Determination Theory）" / "index.html"; source.parent.mkdir()
            source.write_text("<html><title>Self-Determination Theory</title><article><h1>Self-Determination Theory</h1><p>Autonomy, competence and relatedness.</p></article></html>", encoding="utf-8")
            report = ingest_theories([source], home)
            self.assertEqual(report["imported"], 1); self.assertFalse((home / "sources" / source.name).exists())
            repeated = ingest_theories([source], home); self.assertEqual(repeated["merged"], 1); self.assertEqual(len(load_jsonl(home / "sources/registry.jsonl")), 1)
            card_path = next((home / "candidates").glob("*.json")); card = json.loads(card_path.read_text(encoding="utf-8"))
            self.assertEqual(card["verification_status"], "candidate")
            verify_theories("compile", home); card = json.loads(card_path.read_text(encoding="utf-8")); self.assertEqual(card["verification_status"], "candidate")
            card.update({"host_review_status": "completed", "constructs": ["autonomy"], "mechanisms": ["need satisfaction"], "boundary_conditions": ["context"], "foundational_sources": [{"doi": "10.0000/example"}], "verification_notes": "Checked against the registered foundational source."})
            card_path.write_text(json.dumps(card, ensure_ascii=False), encoding="utf-8")
            verify_theories("compile", home); card = json.loads(card_path.read_text(encoding="utf-8")); self.assertEqual(card["verification_status"], "source-located")
            self.assertEqual(promote_theories([card["theory_id"]], False, home)["promoted"], 0)
            self.assertEqual(promote_theories([card["theory_id"]], True, home)["promoted"], 1)
            self.assertEqual(len(search_theories("autonomy need", 8, home, verified_only=True)), 1)

    def test_theory_task_support_is_fail_open_for_empty_or_corrupt_library(self):
        with tempfile.TemporaryDirectory() as td, patch.dict(os.environ, {"CQ_THEORY_HOME": str(Path(td) / "library")}):
            root = ensure_task(Path(td) / "task"); write_json(root / "manifest.json", {"theory_support_mode": "local-only"})
            write_json(root / "00_plan/search_plan.json", {"title_or_idea": "consumer mechanism"})
            home = ensure_library(); (home / "candidates/broken.json").write_text("{broken", encoding="utf-8")
            prepared = safe_task_theory_phase(root, "prepare"); self.assertIn(prepared["status"], {"degraded", "exploratory"})
            compiled = safe_task_theory_phase(root, "compile"); self.assertEqual(compiled["status"], "degraded"); self.assertTrue(compiled["non_blocking"])
            checked = safe_task_theory_phase(root, "validate"); self.assertTrue(checked["valid"]); self.assertTrue(checked["writing_continues"])
            self.assertEqual(checked["writing_fallback_mode"], "mechanism-boundary-neutral-narrative")

    def test_theory_support_cli_never_blocks_task(self):
        with tempfile.TemporaryDirectory() as td, patch.dict(os.environ, {"CQ_THEORY_HOME": str(Path(td) / "library")}):
            root = ensure_task(Path(td) / "task"); write_json(root / "manifest.json", {})
            self.assertEqual(pipeline_main(["theory-support", "--task", str(root), "--mode", "skip"]), 0)
            self.assertEqual(read_json(root / "manifest.json", {})["theory_support_mode"], "skip")
            self.assertEqual(pipeline_main(["theory-recommend", "--task", str(root), "--phase", "validate"]), 0)
            self.assertTrue(read_json(root / "07_logs/theory_validation.json", {})["valid"])

    def test_status_requests_theory_choice_only_at_useful_checkpoint(self):
        with tempfile.TemporaryDirectory() as td, patch.dict(os.environ, {"CQ_THEORY_HOME": str(Path(td) / "library")}):
            root = ensure_task(Path(td) / "task"); record = canonical_record({"title": "Mechanism evidence", "abstract": "A mechanism is studied."})
            write_jsonl(root / "02_corpus/corpus.jsonl", [record])
            write_json(root / "00_plan/search_plan.json", {"approved": True})
            write_json(root / "manifest.json", {"acquisition_mode": "api-search", "corpus_mode": "all", "theory_support_mode": "unselected"})
            (root / "04_analysis/tables/bibliometric_analysis.xlsx").touch(); (root / "05_evidence/cards/card.md").write_text("card", encoding="utf-8")
            (root / "06_review/outline.md").write_text("> status: approved\n\n## 理论机制", encoding="utf-8")
            write_json(root / "06_review/review_brief.json", {"model_package_requested": True})
            write_json(root / "07_logs/gap_validation.json", {"status": "validated"})
            output = io.StringIO()
            with redirect_stdout(output): self.assertEqual(pipeline_main(["status", "--task", str(root)]), 0)
            status = json.loads(output.getvalue())
            self.assertTrue(status["theory_support_needed"]); self.assertEqual(status["checkpoint"], "theory-support-choice"); self.assertTrue(status["requires_user_input"])

    def test_theory_bundle_rejects_path_traversal(self):
        import zipfile
        with tempfile.TemporaryDirectory() as td:
            base = Path(td); bundle = base / "theories.zip"; home = base / "library"
            with zipfile.ZipFile(bundle, "w") as archive:
                archive.writestr("../escape.json", "{}")
                archive.writestr("library.json", '{"schema_version": 1}')
            report = import_library(bundle, home)
            self.assertIn("../escape.json", report["rejected"]); self.assertFalse((base / "escape.json").exists())


if __name__ == "__main__": unittest.main()
